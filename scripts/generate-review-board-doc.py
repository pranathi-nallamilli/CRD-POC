#!/usr/bin/env python3
"""Generate Review Board + CI/CD Integration Word document for Learn & Share session."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level, size, color in [(1, 24, '1A5276'), (2, 18, '2471A3'), (3, 14, '2E86C1')]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(*bytes.fromhex(color))
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_code_block(code, language=''):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    # Add shading
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F4F6F7'
    })
    shading.append(shd)

def add_bullet(text, bold_prefix=''):
    """Add a bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(text):
    """Add a note/callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run('Note: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

# =========================================================
# TITLE PAGE
# =========================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Review Board + CI/CD Integration')
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
r.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Learn & Share Session')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = details.add_run('Pre-Merge Code Review with CI/CD Pipeline Integration\n')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x56, 0x6B, 0x73)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Pranathi Nallamilli\n')
r.font.size = Pt(14)
r.bold = True
r2 = meta.add_run('April 2026\n\n')
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

tech = doc.add_paragraph()
tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tech.add_run('Java Spring Boot  |  Review Board  |  GitHub Actions  |  ArgoCD  |  Kubernetes')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
r.italic = True

doc.add_page_break()

# =========================================================
# TABLE OF CONTENTS (manual)
# =========================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. What is Review Board?',
    '2. Architecture (Local Setup)',
    '3. Setup Steps',
    '4. Demo Workflow (End-to-End)',
    '5. CI/CD Integration',
    '6. Troubleshooting & Lessons Learned',
    '7. Review Board vs GitHub PRs',
    '8. Key Takeaways & Recommendation',
    '9. Quick Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# =========================================================
# 1. WHAT IS REVIEW BOARD?
# =========================================================
doc.add_heading('1. What is Review Board?', level=1)

p = doc.add_paragraph()
p.add_run('Review Board').bold = True
p.add_run(' is a ')
p.add_run('self-hosted code review tool').bold = True
p.add_run(' (not a vulnerability scanner). It provides a web-based interface for reviewing code changes before they are merged, similar to GitHub Pull Request reviews but as a standalone platform.')

doc.add_paragraph()
add_bullet('https://www.reviewboard.org', bold_prefix='Website: ')
add_bullet('Open Source (MIT)', bold_prefix='License: ')
add_bullet('7.0.6', bold_prefix='Latest Version: ')
add_bullet('Python/Django', bold_prefix='Built With: ')

doc.add_paragraph()
doc.add_heading('Feature Comparison', level=3)
add_table(
    ['Feature', 'Review Board', 'GitHub PRs'],
    [
        ['Code diff view', 'Yes', 'Yes'],
        ['Inline comments', 'Yes', 'Yes'],
        ['Approval workflow', '"Ship It!"', '"Approve"'],
        ['CI integration', 'Via API', 'Native (Actions)'],
        ['Self-hosted', 'Required', 'Optional (Enterprise)'],
        ['Pre-commit review', 'Yes (designed for this)', 'Yes (via PRs)'],
        ['Multi-SCM support', 'Git, SVN, Mercurial, Perforce', 'Git only'],
        ['Cost', 'Free (self-hosted infra)', 'Free (public repos)'],
    ],
    col_widths=[2.0, 2.5, 2.5]
)

# =========================================================
# 2. ARCHITECTURE
# =========================================================
doc.add_heading('2. Architecture (Local Setup)', level=1)

p = doc.add_paragraph('The local setup uses 4 containers on a Podman network:')
doc.add_paragraph()

# Architecture diagram as text
add_code_block(
    '  Browser (localhost:8084)\n'
    '      │\n'
    '      ▼\n'
    '  ┌──────────────┐     ┌──────────────────┐\n'
    '  │  Nginx:8084  │────▶│ Review Board:8080│\n'
    '  │ (static files│     │  (Django/Python)  │\n'
    '  │  + proxy)    │     └────────┬─────────┘\n'
    '  └──────────────┘              │\n'
    '                         ┌──────┴──────┐\n'
    '                         │             │\n'
    '                   ┌─────┴───┐  ┌──────┴─────┐\n'
    '                   │Postgres │  │ Memcached  │\n'
    '                   │  :5432  │  │  :11211    │\n'
    '                   └─────────┘  └────────────┘'
)

doc.add_paragraph()
doc.add_heading('Container Details', level=3)
add_table(
    ['Container', 'Image', 'Port', 'Purpose'],
    [
        ['db', 'postgres:15', '5432', 'Database'],
        ['memcached', 'memcached:latest', '11211', 'Caching layer'],
        ['reviewboard', 'beanbag/reviewboard:latest', '8080', 'Review Board application'],
        ['rb-nginx', 'nginx:latest', '8084 → 80', 'Web frontend + static files'],
    ],
    col_widths=[1.3, 2.5, 1.0, 2.2]
)

add_note('All containers run on Podman network "reviewboard-net". Container runtime: Podman v5.6.0 (not Docker).')

# =========================================================
# 3. SETUP STEPS
# =========================================================
doc.add_heading('3. Setup Steps', level=1)

doc.add_heading('Prerequisites', level=2)
add_bullet('Podman v5.6.0+')
add_bullet('Python 3.13+ with pip')
add_bullet('Java 17+, Maven 3.9+')
add_bullet('GitHub account with Personal Access Token (PAT)')

doc.add_heading('Step 1: Create Podman Network', level=2)
add_code_block('podman network create reviewboard-net')

doc.add_heading('Step 2: Start PostgreSQL', level=2)
add_code_block(
    'podman run -d --name db --network reviewboard-net \\\n'
    '  -e POSTGRES_USER=reviewboard \\\n'
    '  -e POSTGRES_PASSWORD=reviewboard123 \\\n'
    '  -e POSTGRES_DB=reviewboard \\\n'
    '  postgres:15'
)

doc.add_heading('Step 3: Start Memcached', level=2)
add_code_block('podman run -d --name memcached --network reviewboard-net memcached:latest')

doc.add_heading('Step 4: Start Review Board', level=2)
add_code_block(
    'podman run -d --name reviewboard --network reviewboard-net \\\n'
    '  -e DATABASE_TYPE=postgresql \\\n'
    '  -e DATABASE_SERVER=db \\\n'
    '  -e DATABASE_USERNAME=reviewboard \\\n'
    '  -e DATABASE_PASSWORD=reviewboard123 \\\n'
    '  -e DATABASE_NAME=reviewboard \\\n'
    '  -e MEMCACHED_SERVER=memcached:11211 \\\n'
    '  -e DOMAIN=localhost:8084 \\\n'
    '  -v ~/reviewboard-site:/site \\\n'
    '  beanbag/reviewboard:latest'
)

doc.add_heading('Step 5: Configure & Start Nginx', level=2)
p = doc.add_paragraph('Create ')
p.add_run('~/reviewboard-config/nginx.conf').bold = True
p.add_run(' with proxy settings, then:')
add_code_block(
    'podman run -d --name rb-nginx --network reviewboard-net \\\n'
    '  -p 8084:80 \\\n'
    '  -v ~/reviewboard-config/nginx.conf:/etc/nginx/conf.d/default.conf:ro \\\n'
    '  -v ~/reviewboard-site/htdocs/static:/var/www/reviewboard/htdocs/static:ro \\\n'
    '  -v ~/reviewboard-site/htdocs/media:/var/www/reviewboard/htdocs/media:ro \\\n'
    '  nginx:latest'
)

add_note('Critical Nginx config: proxy_set_header Host localhost:8084; — must include the port, otherwise Review Board JS AJAX calls fail silently.')

doc.add_heading('Step 6: Configure Review Board', level=2)
add_bullet('Open http://localhost:8084 — login as admin / admin')
add_bullet('Go to Admin → Repositories → Add Repository')
add_bullet('Hosting Service: GitHub, Repository: CRD-POC')
add_bullet('Provide GitHub PAT for API access')

doc.add_heading('Step 7: Install RBTools CLI', level=2)
add_code_block(
    'pip3 install RBTools\n'
    'export PATH="$PATH:$HOME/Library/Python/3.13/bin"'
)

doc.add_heading('Step 8: Configure .reviewboardrc', level=2)
add_code_block(
    'REVIEWBOARD_URL = "http://localhost:8083"\n'
    'REPOSITORY = "CRD-POC"\n'
    'BRANCH = "main"\n'
    'TRACKING_BRANCH = "origin/main"'
)

doc.add_heading('Step 9: Create Users', level=2)
add_table(
    ['User', 'Password', 'Role'],
    [
        ['admin', 'admin', 'Owner — posts review requests'],
        ['reviewer1', 'reviewer1', 'Reviewer — approves with Ship It'],
    ],
    col_widths=[1.5, 1.5, 4.0]
)

doc.add_page_break()

# =========================================================
# 4. DEMO WORKFLOW
# =========================================================
doc.add_heading('4. Demo Workflow (End-to-End)', level=1)

p = doc.add_paragraph('The complete workflow tested and verified:')
doc.add_paragraph()

steps = [
    ('Step 1: Make a Code Change',
     'Edit DemoApplication.java — update the message and version number.\nEdit DemoApplicationTests.java — update test assertion to match.',
     'git checkout feature/test-review\n# Edit files...'),
    ('Step 2: Run Tests Locally',
     'Verify all 4 tests pass before submitting for review.',
     'mvn clean test\n# Tests run: 4, Failures: 0, Errors: 0 — BUILD SUCCESS'),
    ('Step 3: Commit and Push',
     'Commit changes and push the feature branch to GitHub.',
     'git add -A\ngit commit -m "feat: Final demo - update message and version to 1.2.0"\ngit push origin feature/test-review'),
    ('Step 4: Post Review Request',
     'Use rbt CLI to create a review request in Review Board.',
     'rbt post --server http://localhost:8084 \\\n'
     '  --username admin --password admin \\\n'
     '  --summary "Final Demo: Update message and version to 1.2.0" \\\n'
     '  --target-people reviewer1 --publish HEAD\n'
     '# → Review request #6 posted at http://localhost:8084/r/6/'),
    ('Step 5: CI Posts Results to Review Board',
     'CI script runs tests, style checks, Dockerfile check, and posts results via API.',
     'REVIEWBOARD_URL="http://localhost:8084" bash scripts/reviewboard-ci-demo.sh\n'
     '# Tests: PASSED ✅  |  Style: PASSED ✅  |  Dockerfile: FOUND ✅'),
    ('Step 6: Reviewer Approves (Ship It!)',
     'Log in as reviewer1 in browser, navigate to review request, click Review → Ship It!',
     '# In browser: http://localhost:8084/r/6/\n'
     '# Login: reviewer1 / reviewer1\n'
     '# Click: Review → Ship It! → OK'),
    ('Step 7: Close as Submitted',
     'Owner closes the review request indicating the code has been merged.',
     '# In browser: Click Close → Submitted\n'
     '# Or via API:\n'
     'curl -u admin:admin -X PUT "http://localhost:8084/api/review-requests/6/" \\\n'
     '  -d "status=submitted"'),
]

for title, desc, code in steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    add_code_block(code)

doc.add_page_break()

# =========================================================
# 5. CI/CD INTEGRATION
# =========================================================
doc.add_heading('5. CI/CD Integration', level=1)

doc.add_heading('Local CI Script', level=2)
p = doc.add_paragraph()
p.add_run('scripts/reviewboard-ci-demo.sh').bold = True
p.add_run(' simulates what a CI pipeline does:')
add_bullet('Runs mvn clean test (unit tests)')
add_bullet('Checks code style (tab characters in Java files)')
add_bullet('Validates Dockerfile exists')
add_bullet('Posts results to Review Board via REST API')

doc.add_paragraph()
doc.add_heading('API Call to Post CI Results', level=3)
add_code_block(
    'curl -u admin:admin -X POST \\\n'
    '  "${REVIEWBOARD_URL}/api/review-requests/${RR_ID}/reviews/" \\\n'
    '  -d "body_top=CI Results: Tests PASSED, Style PASSED, Dockerfile FOUND" \\\n'
    '  -d "public=true"'
)

doc.add_heading('GitHub Actions Workflow', level=2)
p = doc.add_paragraph()
p.add_run('.github/workflows/review-board-ci.yml').bold = True
p.add_run(' — triggers on pull requests to main:')
add_bullet('Runs unit tests with Maven')
add_bullet('Checks code style')
add_bullet('Posts results as a PR comment using actions/github-script')

doc.add_heading('Integration Flow', level=2)
add_code_block(
    'Developer              Review Board           GitHub Actions        ArgoCD\n'
    '   │                        │                       │                  │\n'
    '   ├─ code change ─────────▶│                       │                  │\n'
    '   ├─ rbt post ────────────▶│ Review Request        │                  │\n'
    '   │                        │◀── CI results ────────┤ (runs tests)     │\n'
    '   │                        │                       │                  │\n'
    '   │             reviewer1 ─┤ Ship It! ✅           │                  │\n'
    '   │                        │                       │                  │\n'
    '   ├─ git merge main ──────▶│ Close: Submitted      │                  │\n'
    '   ├─ git push ────────────────────────────────────▶│ build+push image│\n'
    '   │                        │                       │─────────────────▶│\n'
    '   │                        │                       │       auto-sync  │\n'
    '   │                        │                       │       to K8s     │'
)

doc.add_page_break()

# =========================================================
# 6. TROUBLESHOOTING
# =========================================================
doc.add_heading('6. Troubleshooting & Lessons Learned', level=1)

issues = [
    ('Ship It Not Working in Browser',
     'Nginx Host header was passing "localhost" without the port. Review Board\'s JavaScript AJAX calls targeted port 80 instead of 8084.',
     'Set proxy_set_header Host localhost:8084; and proxy_set_header X-Forwarded-Port 8084; in nginx.conf.'),
    ('Owner Cannot Ship It Own Review',
     'Review Board by design prevents the review request owner from approving their own code. Clicking Ship It as owner has no effect.',
     'Log in as a different user (reviewer1) to approve. This is intentional — you shouldn\'t approve your own code.'),
    ('CSRF 403 Error on Login',
     'Domain mismatch between browser URL and Review Board\'s configured site domain.',
     'Add CSRF_TRUSTED_ORIGINS = ["http://localhost:8084", "http://localhost:8083"] to settings_local.py.'),
    ('rbt post "Connection Refused"',
     '.reviewboardrc pointed to Nginx port (8084), but API redirect URLs pointed to port 80 before the Nginx fix.',
     'Use direct backend port 8083 in .reviewboardrc, or fix the Nginx Host header (recommended).'),
    ('502 Bad Gateway',
     'Review Board container had stopped.',
     'podman start db memcached reviewboard rb-nginx'),
]

for title, cause, fix in issues:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run('Cause: ').bold = True
    p.add_run(cause)
    p2 = doc.add_paragraph()
    p2.add_run('Fix: ').bold = True
    p2.add_run(fix)
    doc.add_paragraph()

doc.add_page_break()

# =========================================================
# 7. REVIEW BOARD vs GITHUB PRs
# =========================================================
doc.add_heading('7. Review Board vs GitHub PRs', level=1)

add_table(
    ['Aspect', 'Review Board', 'GitHub PRs'],
    [
        ['Best for', 'Pre-commit review, large enterprises', 'GitHub-native workflows'],
        ['Setup', 'Self-hosted (4 containers)', 'Zero setup'],
        ['SCM support', 'Git, SVN, Mercurial, Perforce', 'Git only'],
        ['CI integration', 'Via API (manual)', 'Native (GitHub Actions)'],
        ['Learning curve', 'Moderate (new tool + CLI)', 'Low (already using GitHub)'],
        ['Team adoption', 'Requires training', 'Already familiar'],
        ['Maintenance', 'DB, containers, updates', 'Managed by GitHub'],
        ['Cost', 'Free + infrastructure cost', 'Free (public repos)'],
    ],
    col_widths=[1.5, 2.75, 2.75]
)

doc.add_heading('When Review Board Makes Sense', level=2)
add_bullet('Teams using multiple SCMs (Git + SVN + Perforce)')
add_bullet('Organizations requiring self-hosted code review (compliance)')
add_bullet('Pre-commit review workflows (before code is pushed)')
add_bullet('Integration with legacy systems not on GitHub')

doc.add_heading('When GitHub PRs Are Sufficient', level=2)
add_bullet('Teams already using GitHub for source control')
add_bullet('CI/CD already built on GitHub Actions')
add_bullet('Branch protection rules enforce review requirements')
add_bullet('No need for multi-SCM support')

doc.add_page_break()

# =========================================================
# 8. KEY TAKEAWAYS
# =========================================================
doc.add_heading('8. Key Takeaways & Recommendation', level=1)

doc.add_heading('What We Learned', level=2)
add_bullet('Review Board is a mature, self-hosted code review tool with a rich API')
add_bullet('It integrates with CI via REST API — CI results can be posted as review comments')
add_bullet('The "Ship It!" workflow enforces that someone other than the author approves')
add_bullet('Local setup requires 4 containers (PostgreSQL, Memcached, Review Board, Nginx)')
add_bullet('RBTools CLI (rbt) is used to post review requests from the command line')
add_bullet('Nginx proxy configuration is critical — Host header must include the port')

doc.add_heading('Recommendation', level=2)
p = doc.add_paragraph()
p.add_run('For teams already on GitHub, ').font.size = Pt(12)
r = p.add_run('GitHub PRs with branch protection rules provide the same code review workflow ')
r.font.size = Pt(12)
r.bold = True
r2 = p.add_run('without the overhead of maintaining a separate tool.')
r2.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph('Review Board adds value when you need to:')
add_bullet('Support multiple SCMs (Git + SVN + Perforce)')
add_bullet('Have self-hosted code review for compliance reasons')
add_bullet('Use pre-commit review workflows')
add_bullet('Integrate with systems outside the GitHub ecosystem')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Key Question for the Team: ')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
p.add_run('"What specific problem are we trying to solve that GitHub PRs don\'t already handle?"')

doc.add_page_break()

# =========================================================
# 9. QUICK REFERENCE
# =========================================================
doc.add_heading('9. Quick Reference', level=1)

doc.add_heading('Start/Stop Review Board', level=2)
add_code_block(
    '# Start all containers\n'
    'podman start db memcached reviewboard rb-nginx\n\n'
    '# Stop all containers\n'
    'podman stop rb-nginx reviewboard memcached db'
)

doc.add_heading('Post a Review Request', level=2)
add_code_block(
    'export PATH="$PATH:$HOME/Library/Python/3.13/bin"\n'
    'rbt post --server http://localhost:8084 \\\n'
    '  --username admin --password admin \\\n'
    '  --target-people reviewer1 \\\n'
    '  --publish HEAD'
)

doc.add_heading('Run CI Demo', level=2)
add_code_block('REVIEWBOARD_URL="http://localhost:8084" bash scripts/reviewboard-ci-demo.sh')

doc.add_heading('Credentials', level=2)
add_table(
    ['User', 'Password', 'Role'],
    [
        ['admin', 'admin', 'Owner (posts review requests)'],
        ['reviewer1', 'reviewer1', 'Reviewer (approves with Ship It)'],
    ],
    col_widths=[1.5, 1.5, 4.0]
)

doc.add_heading('URLs', level=2)
add_table(
    ['Service', 'URL'],
    [
        ['Review Board (browser)', 'http://localhost:8084'],
        ['Review Board (API/CLI)', 'http://localhost:8083'],
        ['Review Board Admin', 'http://localhost:8084/admin/'],
        ['GitHub Repository', 'https://github.com/pranathi-nallamilli/CRD-POC'],
    ],
    col_widths=[2.5, 4.5]
)

# -- Save --
output_path = os.path.expanduser('~/Documents/pranathi-learning/CRD-POC/Review-Board-CI-CD-Demo.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
