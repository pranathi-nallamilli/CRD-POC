#!/usr/bin/env python3
"""
Generate colorful CI/CD architecture diagrams and build a Word document.
Inspired by the ELK Stack architecture diagram style.
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import os

# ─── Color Palette ───
COLORS = {
    'github':       '#24292e',   # dark gray
    'github_text':  '#ffffff',
    'actions':      '#2088FF',   # GitHub Actions blue
    'actions_text': '#ffffff',
    'docker':       '#2496ED',   # Docker blue
    'docker_text':  '#ffffff',
    'argocd':       '#EF7B4D',   # ArgoCD orange
    'argocd_text':  '#ffffff',
    'k8s':          '#326CE5',   # Kubernetes blue
    'k8s_text':     '#ffffff',
    'pod_green':    '#4CAF50',   # Green for running pods
    'pod_text':     '#ffffff',
    'cosign':       '#7B1FA2',   # Purple for security
    'cosign_text':  '#ffffff',
    'kyverno':      '#D32F2F',   # Red for admission control
    'kyverno_text': '#ffffff',
    'updater':      '#FF9800',   # Orange for Image Updater
    'updater_text': '#ffffff',
    'trivy':        '#1DE9B6',   # Teal for scanning
    'trivy_text':   '#000000',
    'job_build':    '#43A047',   # Green
    'job_docker':   '#1E88E5',   # Blue
    'job_manifest': '#FB8C00',   # Orange
    'job_notify':   '#8E24AA',   # Purple
    'job_text':     '#ffffff',
    'cluster_bg':   '#E3F2FD',   # Light blue background
    'cluster_border': '#1565C0',
    'service':      '#F48FB1',   # Pink
    'service_text': '#000000',
    'arrow':        '#546E7A',
    'bg':           '#FAFAFA',
}


def draw_rounded_box(ax, x, y, w, h, color, label, sublabel=None, fontsize=11, sublabel_size=8):
    """Draw a colored rounded rectangle with centered text."""
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.02",
                         facecolor=color, edgecolor='white',
                         linewidth=2, zorder=3)
    ax.add_patch(box)
    
    text_color = '#ffffff' if sum(int(color.lstrip('#')[i:i+2], 16) for i in (0,2,4)) < 400 else '#000000'
    
    if sublabel:
        ax.text(x + w/2, y + h*0.62, label, ha='center', va='center',
                fontsize=fontsize, fontweight='bold', color=text_color, zorder=4)
        ax.text(x + w/2, y + h*0.30, sublabel, ha='center', va='center',
                fontsize=sublabel_size, color=text_color, alpha=0.9, zorder=4)
    else:
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=fontsize, fontweight='bold', color=text_color, zorder=4)


def draw_arrow(ax, x1, y1, x2, y2, label=None, color=None):
    """Draw an arrow between two points with optional label."""
    if color is None:
        color = COLORS['arrow']
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=2.5,
                               connectionstyle='arc3,rad=0'),
                zorder=2)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my + 0.15, label, ha='center', va='bottom',
                fontsize=7.5, color=color, fontweight='bold', zorder=5,
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white', edgecolor='none', alpha=0.85))


# ═══════════════════════════════════════════════════════════════
#  APPROACH 1 DIAGRAM
# ═══════════════════════════════════════════════════════════════
def generate_approach1():
    fig, ax = plt.subplots(1, 1, figsize=(16, 11))
    ax.set_xlim(-0.5, 15.5)
    ax.set_ylim(-1, 11)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor(COLORS['bg'])
    
    # Title
    ax.text(7.5, 10.5, 'Approach 1: Pipeline-Driven Manifest Update',
            ha='center', va='center', fontsize=18, fontweight='bold', color='#1565C0')
    ax.text(7.5, 10.0, 'GitHub Actions (4 Jobs) → ArgoCD → Kubernetes',
            ha='center', va='center', fontsize=11, color='#666666')

    # ─── GitHub Repository ───
    repo_bg = FancyBboxPatch((1, 8.2), 13, 1.3, boxstyle="round,pad=0.05",
                              facecolor='#F5F5F5', edgecolor=COLORS['github'], linewidth=2, linestyle='--', zorder=1)
    ax.add_patch(repo_bg)
    ax.text(7.5, 9.25, 'GITHUB REPOSITORY', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['github'])
    
    draw_rounded_box(ax, 1.5, 8.35, 2.8, 0.7, '#455A64', 'Source Code', 'Java 17 / Maven', 9, 7)
    draw_rounded_box(ax, 5.5, 8.35, 2.8, 0.7, '#455A64', 'Dockerfile', 'Multi-stage', 9, 7)
    draw_rounded_box(ax, 9.5, 8.35, 2.8, 0.7, '#455A64', 'K8s Manifests', 'GitOps', 9, 7)

    # Arrow: git push
    draw_arrow(ax, 7.5, 8.2, 7.5, 7.5, 'git push', '#24292e')

    # ─── GitHub Actions Pipeline ───
    pipe_bg = FancyBboxPatch((0.5, 5.5), 14, 1.9, boxstyle="round,pad=0.05",
                              facecolor='#E8EAF6', edgecolor=COLORS['actions'], linewidth=2, zorder=1)
    ax.add_patch(pipe_bg)
    ax.text(7.5, 7.15, 'GITHUB ACTIONS PIPELINE (4 Jobs)', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['actions'])

    # 4 Jobs
    draw_rounded_box(ax, 1, 5.7, 2.5, 1.0, COLORS['job_build'], 'Job 1', 'Build & Test', 10, 8)
    draw_rounded_box(ax, 4.2, 5.7, 2.5, 1.0, COLORS['job_docker'], 'Job 2', 'Docker Push', 10, 8)
    draw_rounded_box(ax, 7.4, 5.7, 2.5, 1.0, COLORS['job_manifest'], 'Job 3', 'Update Manifest', 10, 8)
    draw_rounded_box(ax, 10.6, 5.7, 2.5, 1.0, COLORS['job_notify'], 'Job 4', 'Notify', 10, 8)
    
    # Job arrows
    draw_arrow(ax, 3.5, 6.2, 4.2, 6.2, color=COLORS['arrow'])
    draw_arrow(ax, 6.7, 6.2, 7.4, 6.2, color=COLORS['arrow'])
    draw_arrow(ax, 9.9, 6.2, 10.6, 6.2, color=COLORS['arrow'])

    # ─── Docker Hub ───
    draw_rounded_box(ax, 2.5, 3.8, 3, 1.0, COLORS['docker'], 'Docker Hub', 'Image: main-<sha>', 10, 8)
    draw_arrow(ax, 5.45, 5.7, 4.0, 4.8, 'Push Image', COLORS['docker'])

    # ─── Git Repo (manifest updated) ───
    draw_rounded_box(ax, 8.5, 3.8, 3.5, 1.0, '#FF7043', 'Git Repo', 'Updated deployment.yaml', 10, 8)
    draw_arrow(ax, 8.65, 5.7, 10.25, 4.8, 'git push manifest', '#FF7043')

    # ─── ArgoCD ───
    draw_rounded_box(ax, 8.5, 2.2, 3.5, 1.0, COLORS['argocd'], 'ArgoCD', 'Polls Git (3 min) → Sync', 10, 8)
    draw_arrow(ax, 10.25, 3.8, 10.25, 3.2, 'Detects change', COLORS['argocd'])

    # ─── Kubernetes Cluster ───
    k8s_bg = FancyBboxPatch((1, -0.5), 13, 2.3, boxstyle="round,pad=0.05",
                             facecolor=COLORS['cluster_bg'], edgecolor=COLORS['cluster_border'],
                             linewidth=2.5, linestyle='--', zorder=1)
    ax.add_patch(k8s_bg)
    ax.text(7.5, 1.55, 'KUBERNETES CLUSTER (kind - 3 Nodes)', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['cluster_border'])

    # 3 Pods
    draw_rounded_box(ax, 2, 0.1, 2, 0.8, COLORS['pod_green'], 'Pod 1', 'Running ✓', 9, 7)
    draw_rounded_box(ax, 5, 0.1, 2, 0.8, COLORS['pod_green'], 'Pod 2', 'Running ✓', 9, 7)
    draw_rounded_box(ax, 8, 0.1, 2, 0.8, COLORS['pod_green'], 'Pod 3', 'Running ✓', 9, 7)
    
    # Service
    draw_rounded_box(ax, 10.8, 0.1, 2.5, 0.8, COLORS['service'], 'Service', 'ClusterIP :80→8080', 9, 7)

    # Arrow: ArgoCD to K8s
    draw_arrow(ax, 10.25, 2.2, 7.5, 1.8, 'Auto Sync', COLORS['k8s'])

    # ─── Legend ───
    legend_items = [
        (COLORS['job_build'], 'Build & Test'),
        (COLORS['job_docker'], 'Docker Build'),
        (COLORS['job_manifest'], 'Manifest Update'),
        (COLORS['job_notify'], 'Notification'),
        (COLORS['argocd'], 'ArgoCD (GitOps)'),
        (COLORS['pod_green'], 'Running Pod'),
    ]
    for i, (color, label) in enumerate(legend_items):
        bx = 0.2 + (i % 3) * 4.5
        by = -0.85
        box = FancyBboxPatch((bx, by), 0.3, 0.2, boxstyle="round,pad=0.02",
                              facecolor=color, edgecolor='white', linewidth=1, zorder=3)
        ax.add_patch(box)
        ax.text(bx + 0.45, by + 0.1, label, va='center', fontsize=8, color='#333333')

    # ─── Warning box ───
    warn_bg = FancyBboxPatch((1, -0.85), 13, 0.0, boxstyle="round,pad=0.02",
                              facecolor='#FFF3E0', edgecolor='#FF9800', linewidth=0, zorder=0)
    ax.add_patch(warn_bg)

    plt.tight_layout()
    path = os.path.join(os.path.dirname(__file__), 'approach1_architecture.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=COLORS['bg'])
    plt.close()
    print(f"✅ Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
#  APPROACH 2 DIAGRAM
# ═══════════════════════════════════════════════════════════════
def generate_approach2():
    fig, ax = plt.subplots(1, 1, figsize=(16, 13))
    ax.set_xlim(-0.5, 15.5)
    ax.set_ylim(-2, 12.5)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor(COLORS['bg'])

    # Title
    ax.text(7.5, 12.0, 'Approach 2: ArgoCD Image Updater + Cosign + Kyverno',
            ha='center', va='center', fontsize=18, fontweight='bold', color='#1565C0')
    ax.text(7.5, 11.5, 'Production-Grade Pipeline (2 Jobs) with Image Signing & Admission Control',
            ha='center', va='center', fontsize=11, color='#666666')

    # ─── GitHub Repository ───
    repo_bg = FancyBboxPatch((1, 9.5), 13, 1.3, boxstyle="round,pad=0.05",
                              facecolor='#F5F5F5', edgecolor=COLORS['github'], linewidth=2, linestyle='--', zorder=1)
    ax.add_patch(repo_bg)
    ax.text(7.5, 10.55, 'GITHUB REPOSITORY', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['github'])
    
    draw_rounded_box(ax, 1.5, 9.65, 2.8, 0.7, '#455A64', 'Source Code', 'Java 17 / Maven', 9, 7)
    draw_rounded_box(ax, 5.5, 9.65, 2.8, 0.7, '#455A64', 'Dockerfile', 'Multi-stage', 9, 7)
    draw_rounded_box(ax, 9.5, 9.65, 2.8, 0.7, '#455A64', 'K8s Manifests', 'Kustomize', 9, 7)

    # Arrow: git push
    draw_arrow(ax, 7.5, 9.5, 7.5, 8.9, 'git push', '#24292e')

    # ─── GitHub Actions Pipeline (2 Jobs) ───
    pipe_bg = FancyBboxPatch((1, 7.0), 13, 1.8, boxstyle="round,pad=0.05",
                              facecolor='#E8EAF6', edgecolor=COLORS['actions'], linewidth=2, zorder=1)
    ax.add_patch(pipe_bg)
    ax.text(7.5, 8.55, 'GITHUB ACTIONS PIPELINE (2 Jobs — Read-Only!)', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['actions'])

    # 2 Jobs
    draw_rounded_box(ax, 1.5, 7.2, 3.5, 1.0, COLORS['job_build'], 'Job 1: Build & Test', 'Maven + Unit Tests', 10, 8)
    draw_rounded_box(ax, 6.5, 7.2, 6.5, 1.0, COLORS['job_docker'], 'Job 2: Docker Build + Push + Trivy Scan + Cosign Sign', '', 9.5, 8)
    
    # Arrow between jobs
    draw_arrow(ax, 5.0, 7.7, 6.5, 7.7, color=COLORS['arrow'])

    # ─── No Job 3/4 callout ───
    ax.text(14.2, 7.7, 'No Job 3!\nNo Job 4!', ha='center', va='center',
            fontsize=8, fontweight='bold', color='#D32F2F',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFEBEE', edgecolor='#D32F2F', linewidth=1.5))

    # Arrow to Docker Hub
    draw_arrow(ax, 9.75, 7.2, 9.75, 6.5, 'Push Image\n+ Signature', COLORS['docker'])

    # ─── Docker Hub ───
    draw_rounded_box(ax, 7.5, 5.3, 4.5, 1.1, COLORS['docker'], 'Docker Hub', 'image:main-<sha> + Cosign Signature', 11, 8)

    # ─── Two paths from Docker Hub ───
    # Left: Image Updater
    draw_arrow(ax, 8.5, 5.3, 4.0, 4.5, 'Watches for\nnew tags', COLORS['updater'])
    # Right: ArgoCD
    draw_arrow(ax, 10.5, 5.3, 12.0, 4.5, '', COLORS['argocd'])

    # ─── ArgoCD Image Updater ───
    draw_rounded_box(ax, 1.5, 3.3, 4.5, 1.1, COLORS['updater'], 'ArgoCD Image Updater', 'Detects new tag → Commits to Git', 10, 8)

    # ─── ArgoCD ───
    draw_rounded_box(ax, 9.5, 3.3, 4.5, 1.1, COLORS['argocd'], 'ArgoCD', 'Monitors Git → Auto-Sync → Self-Heal', 10, 8)

    # Arrow: Image Updater → Git → ArgoCD
    draw_arrow(ax, 6.0, 3.85, 9.5, 3.85, 'Updates Git manifest', '#546E7A')

    # Arrow: ArgoCD → K8s
    draw_arrow(ax, 11.75, 3.3, 11.75, 2.6, 'Deploy', COLORS['k8s'])

    # ─── Kubernetes Cluster ───
    k8s_bg = FancyBboxPatch((0.5, -1.5), 14, 3.8, boxstyle="round,pad=0.05",
                             facecolor=COLORS['cluster_bg'], edgecolor=COLORS['cluster_border'],
                             linewidth=2.5, linestyle='--', zorder=1)
    ax.add_patch(k8s_bg)
    ax.text(7.5, 2.05, 'KUBERNETES CLUSTER (kind / EKS / AKS / GKE)', ha='center', va='center',
            fontsize=11, fontweight='bold', color=COLORS['cluster_border'])

    # ─── Kyverno ───
    kyverno_bg = FancyBboxPatch((1.5, 0.7), 12, 1.0, boxstyle="round,pad=0.03",
                                 facecolor='#FFEBEE', edgecolor=COLORS['kyverno'], linewidth=2, zorder=2)
    ax.add_patch(kyverno_bg)
    ax.text(7.5, 1.35, 'Kyverno — Admission Controller', ha='center', va='center',
            fontsize=10, fontweight='bold', color=COLORS['kyverno'])
    ax.text(7.5, 0.95, '"Is this image signed by our Cosign key?"    [YES] Signed = Allow    [NO] Unsigned = BLOCK',
            ha='center', va='center', fontsize=8, color='#555555')

    # 3 Pods
    draw_rounded_box(ax, 1.5, -0.6, 2.5, 0.9, COLORS['pod_green'], 'Pod 1', 'Verified', 10, 8)
    draw_rounded_box(ax, 5, -0.6, 2.5, 0.9, COLORS['pod_green'], 'Pod 2', 'Verified', 10, 8)
    draw_rounded_box(ax, 8.5, -0.6, 2.5, 0.9, COLORS['pod_green'], 'Pod 3', 'Verified', 10, 8)
    
    # Service
    draw_rounded_box(ax, 11.5, -0.6, 2.5, 0.9, COLORS['service'], 'Service', 'ClusterIP', 9, 7)

    # ─── Security Layers Box ───
    sec_bg = FancyBboxPatch((0.5, -1.95), 14, 0.35, boxstyle="round,pad=0.02",
                             facecolor='#E8F5E9', edgecolor='#4CAF50', linewidth=1.5, zorder=1)
    ax.add_patch(sec_bg)
    ax.text(7.5, -1.78, '5 Security Layers:  Trivy Scan → Cosign Sign → Docker Hub IAM → Kyverno Verify → ArgoCD Self-Heal',
            ha='center', va='center', fontsize=8, fontweight='bold', color='#2E7D32')

    # ─── Legend ───
    legend_items = [
        (COLORS['job_build'], 'Build & Test'),
        (COLORS['job_docker'], 'Docker + Sign'),
        (COLORS['updater'], 'Image Updater'),
        (COLORS['argocd'], 'ArgoCD (GitOps)'),
        (COLORS['kyverno'], 'Kyverno (Admission)'),
        (COLORS['pod_green'], 'Verified Pod'),
    ]
    for i, (color, label) in enumerate(legend_items):
        bx = 0.5 + (i % 6) * 2.4
        by = -1.45
        box = FancyBboxPatch((bx, by), 0.25, 0.18, boxstyle="round,pad=0.02",
                              facecolor=color, edgecolor='white', linewidth=1, zorder=3)
        ax.add_patch(box)
        ax.text(bx + 0.38, by + 0.09, label, va='center', fontsize=7, color='#333333')

    plt.tight_layout()
    path = os.path.join(os.path.dirname(__file__), 'approach2_architecture.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=COLORS['bg'])
    plt.close()
    print(f"✅ Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
#  DATA FLOW COMPARISON DIAGRAM
# ═══════════════════════════════════════════════════════════════
def generate_comparison():
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 6))
    fig.patch.set_facecolor(COLORS['bg'])

    for ax, title, color_main, steps, step_colors in [
        (ax1, 'Approach 1: Data Flow (4 Jobs — Pipeline writes to Git)',
         '#FF7043',
         ['Developer\ngit push', 'Job 1\nBuild & Test', 'Job 2\nDocker Push', 'Job 3\nUpdate Manifest\n(git push!)', 'Job 4\nNotify', 'ArgoCD\nPolls Git', 'K8s\n3 Pods'],
         ['#455A64', COLORS['job_build'], COLORS['job_docker'], '#FF7043', COLORS['job_notify'], COLORS['argocd'], COLORS['pod_green']]),
        (ax2, 'Approach 2: Data Flow (2 Jobs — Pipeline is read-only)',
         '#1565C0',
         ['Developer\ngit push', 'Job 1\nBuild & Test', 'Job 2\nDocker+Sign', 'Image Updater\nDetects tag', 'ArgoCD\nSyncs', 'Kyverno\nVerifies sig', 'K8s\n3 Pods'],
         ['#455A64', COLORS['job_build'], COLORS['job_docker'], COLORS['updater'], COLORS['argocd'], COLORS['kyverno'], COLORS['pod_green']]),
    ]:
        ax.set_xlim(-0.5, 14.5)
        ax.set_ylim(-0.5, 2.5)
        ax.axis('off')
        ax.set_title(title, fontsize=11, fontweight='bold', color=color_main, pad=10)

        n = len(steps)
        spacing = 14.0 / n
        for i, (step, clr) in enumerate(zip(steps, step_colors)):
            x = i * spacing + 0.3
            box = FancyBboxPatch((x, 0.2), spacing * 0.8, 1.5,
                                  boxstyle="round,pad=0.05", facecolor=clr,
                                  edgecolor='white', linewidth=2, zorder=3)
            ax.add_patch(box)
            tc = '#ffffff' if sum(int(clr.lstrip('#')[j:j+2], 16) for j in (0,2,4)) < 400 else '#000000'
            ax.text(x + spacing * 0.4, 0.95, step, ha='center', va='center',
                    fontsize=8, fontweight='bold', color=tc, zorder=4)
            
            if i < n - 1:
                ax.annotate('', xy=(x + spacing * 0.85, 0.95),
                           xytext=(x + spacing * 0.8 + spacing * 0.15, 0.95),
                           arrowprops=dict(arrowstyle='->', color='#78909C', lw=2))

    plt.tight_layout(pad=1.5)
    path = os.path.join(os.path.dirname(__file__), 'comparison_flow.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=COLORS['bg'])
    plt.close()
    print(f"✅ Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
#  BUILD DOCX
# ═══════════════════════════════════════════════════════════════
def build_docx(img1, img2, img3):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ─── Title ───
    title = doc.add_heading('CI/CD Pipeline with GitOps — Two Approaches', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Java Spring Boot | GitHub Actions | ArgoCD | Kubernetes')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    
    # ─── Links ───
    p = doc.add_paragraph()
    run = p.add_run('Repository: ')
    run.bold = True
    p.add_run('https://github.com/pranathi-nallamilli/CRD-POC')
    p = doc.add_paragraph()
    run = p.add_run('Docker Hub: ')
    run.bold = True
    p.add_run('https://hub.docker.com/r/pranathinallamilli/java-demo-app')

    doc.add_page_break()

    # ─── Application Overview ───
    doc.add_heading('Application Overview', level=1)
    table = doc.add_table(rows=8, cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ('Component', 'Detail'),
        ('Language', 'Java 17 (Temurin)'),
        ('Framework', 'Spring Boot 3.2.0'),
        ('Build Tool', 'Maven 3.9+'),
        ('Endpoints', 'GET / (message), GET /health, GET /info'),
        ('Replicas', '3 pods (HA with rolling updates)'),
        ('Resources', '256Mi–512Mi memory, 250m–500m CPU per pod'),
        ('Cluster', 'kind (3-node: 1 control-plane + 2 workers)'),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # ─── Why kind ───
    doc.add_heading('Why kind Instead of Minikube?', level=1)
    doc.add_paragraph(
        'We initially tried Minikube, but ArgoCD\'s repo-server kept crashing '
        '(CrashLoopBackOff) because it couldn\'t reach GitHub to clone the repo. '
        'We switched to kind and everything worked immediately.'
    )
    doc.add_paragraph(
        'Root Cause: Minikube runs inside a VM (box inside a box). DNS inside the VM '
        'couldn\'t resolve external URLs like github.com on macOS. '
        'kind runs containers directly on the host network — DNS works out of the box.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Simple analogy: ')
    run.bold = True
    p.add_run('Minikube = phone call through two walls (signal lost). '
              'kind = phone call in the same room (clear connection).')

    doc.add_page_break()

    # ─── Approach 1 ───
    doc.add_heading('Approach 1: Pipeline-Driven Manifest Update', level=1)
    doc.add_heading('Architecture Diagram', level=2)
    doc.add_picture(img1, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('Pipeline Flow', level=2)
    doc.add_paragraph('Job 1: Build & Test — Maven clean test → Build JAR → Upload artifact')
    doc.add_paragraph('Job 2: Docker Build & Push — Multi-stage build → Tag main-<sha> → Push to Docker Hub → Trivy scan (fails on CRITICAL/HIGH CVEs) → Cosign sign')
    doc.add_paragraph('Job 3: Update Manifest — sed updates image tag in deployment.yaml → git push back to repo')
    doc.add_paragraph('Job 4: Notify — Report status via GitHub Script API')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Key config: ')
    run.bold = True
    p.add_run('paths-ignore: k8s/deployment.yaml prevents Job 3\'s git push from re-triggering the pipeline.')

    doc.add_heading('Limitations', level=2)
    limitations = [
        'GIT_TOKEN required — extra secret to manage, security risk if leaked',
        'Infinite loop risk — needs paths-ignore workaround; fragile',
        'Pipeline writes to Git — violates GitOps principle (CI shouldn\'t modify source of truth)',
        'No image signing — can\'t verify who built the image',
        'No admission control — K8s accepts ANY image, even malicious ones',
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    doc.add_page_break()

    # ─── Approach 2 ───
    doc.add_heading('Approach 2: ArgoCD Image Updater + Cosign (Production)', level=1)
    doc.add_heading('Architecture Diagram', level=2)
    doc.add_picture(img2, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('How It Works', level=2)
    steps = [
        'Developer pushes code → GitHub Actions runs 2 jobs (Build+Test → Docker+Sign)',
        'Pipeline builds Docker image, pushes to Docker Hub, signs with Cosign — DONE',
        'ArgoCD Image Updater (inside cluster) watches Docker Hub for new tags',
        'Image Updater detects new tag → commits updated manifest to Git',
        'ArgoCD detects Git change → syncs to Kubernetes',
        'Kyverno verifies Cosign signature before allowing pod to start',
        'Only signed images run — unsigned images are BLOCKED',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph()
    doc.add_heading('Why Approach 2 is Better for Production', level=2)
    benefits = [
        ('No GIT_TOKEN needed', 'Pipeline doesn\'t write to Git; eliminates a security secret'),
        ('No infinite loop risk', 'Pipeline never modifies manifests; no paths-ignore needed'),
        ('Separation of concerns', 'CI builds images, ArgoCD handles deployment entirely'),
        ('Image verification', 'Cosign ensures only YOUR pipeline\'s images can be deployed'),
        ('Admission control', 'Kyverno blocks all unsigned/untrusted images at the K8s API level'),
        ('Simpler pipeline', '2 jobs instead of 4; less code to maintain'),
    ]
    table = doc.add_table(rows=len(benefits)+1, cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = 'Benefit'
    table.rows[0].cells[1].text = 'Explanation'
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, (benefit, explanation) in enumerate(benefits, 1):
        table.rows[i].cells[0].text = benefit
        table.rows[i].cells[1].text = explanation

    doc.add_page_break()

    # ─── Cosign + Kyverno ───
    doc.add_heading('How Cosign + Kyverno Eliminate Malicious Images', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('The Problem: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    p.add_run('Without signing, if an attacker gains Docker Hub access, they can push a malicious '
              'image with a legitimate-looking tag. ArgoCD will deploy it — no questions asked.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cosign = Digital Signature for Docker Images')
    run.bold = True
    run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
    doc.add_paragraph('Your CI pipeline has a PRIVATE KEY (in GitHub Secrets). '
                      'After building the image, the pipeline SIGNS it. '
                      'Even if an attacker pushes a fake image, they CAN\'T sign it — they don\'t have the key.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Kyverno = Security Guard at the Kubernetes Door')
    run.bold = True
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    doc.add_paragraph('Before ANY pod starts, Kyverno checks: "Is this image signed by our key?" '
                      'Signed → Allow. Unsigned → BLOCK. '
                      'Even if Docker Hub is compromised, unsigned images can NEVER run in your cluster.')

    doc.add_paragraph()
    doc.add_heading('Attack Scenario Comparison', level=2)
    attacks = [
        ('Attack', 'Without Cosign+Kyverno', 'With Cosign+Kyverno'),
        ('Attacker pushes malicious image', '❌ Deployed to production', '✅ Blocked (no signature)'),
        ('Attacker replaces latest tag', '❌ Deployed on next sync', '✅ Blocked (signature mismatch)'),
        ('Compromised CI in another repo', '❌ Could deploy bad images', '✅ Blocked (wrong signing key)'),
        ('Supply chain attack on base image', '❌ Silently deployed', '✅ Blocked (final image unsigned)'),
    ]
    table = doc.add_table(rows=len(attacks), cols=3, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(attacks):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_page_break()

    # ─── Comparison Flow ───
    doc.add_heading('Pipeline Data Flow Comparison', level=1)
    doc.add_picture(img3, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ─── Security Architecture ───
    doc.add_heading('Security Architecture: 5 Layers of Protection', level=1)
    layers = [
        ('Layer 1: Build', 'Trivy', 'Scans image for CVEs — fails pipeline on CRITICAL/HIGH (exit-code: 1)', 'During CI (after push, before sign)'),
        ('Layer 2: Sign', 'Cosign', 'Creates cryptographic signature', 'After Docker push'),
        ('Layer 3: Registry', 'Docker Hub IAM', 'Controls who can push/pull', 'Always'),
        ('Layer 4: Verify', 'Kyverno', 'Rejects unsigned images at K8s API', 'Before pod creation'),
        ('Layer 5: Runtime', 'ArgoCD self-heal', 'Reverts manual changes', 'Continuous'),
    ]
    table = doc.add_table(rows=len(layers)+1, cols=4, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Layer', 'Tool', 'What It Does', 'When']):
        table.rows[0].cells[j].text = h
        for paragraph in table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row_data in enumerate(layers, 1):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Result: Even if Docker Hub is fully compromised, unsigned images cannot run in your cluster.')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_page_break()

    # ─── Cost Comparison ───
    doc.add_heading('Pipeline & Cost Comparison', level=1)
    metrics = [
        ('Metric', 'Approach 1', 'Approach 2'),
        ('Pipeline jobs', '4', '2'),
        ('Pipeline time', '~4-5 min', '~3-4 min'),
        ('Git write access', 'Yes (risky)', 'No'),
        ('Image signing', 'No', 'Yes (Cosign)'),
        ('Admission control', 'No', 'Yes (Kyverno)'),
        ('GitHub Actions cost', '$0 (free tier)', '$0 (free tier)'),
        ('K8s cluster (prod)', '~$73/mo (EKS/AKS)', '~$73/mo (EKS/AKS)'),
        ('Extra tools cost', '$0', '$0 (all open source)'),
    ]
    table = doc.add_table(rows=len(metrics), cols=3, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(metrics):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # ─── Why Trivy After Docker Push ───
    doc.add_heading('Why Is Trivy Scan After Docker Push?', level=1)
    doc.add_paragraph(
        'This is a deliberate design choice. docker buildx builds multi-architecture images (AMD64 + ARM64) '
        'and pushes directly to the registry — there is no local image to scan. Trivy needs the final '
        'multi-arch manifest on the registry to scan accurately.'
    )
    doc.add_paragraph(
        'The safety net: If Trivy finds CRITICAL/HIGH CVEs → pipeline fails → image is NEVER signed → '
        'Kyverno BLOCKS it from deploying. So even though the image is on Docker Hub, it can never reach your pods.'
    )
    trivy_flow = [
        ('Step', 'Action', 'What Happens on Failure'),
        ('1. Build & Push', 'docker buildx builds multi-arch image, pushes to Docker Hub', 'Pipeline fails — no image published'),
        ('2. Trivy Scan', 'Scans pushed image for CRITICAL/HIGH CVEs', 'Pipeline fails — image stays UNSIGNED'),
        ('3. Cosign Sign', 'Signs image with private key', 'Without signature, Kyverno blocks deployment'),
    ]
    table = doc.add_table(rows=len(trivy_flow), cols=3, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(trivy_flow):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_page_break()

    # ─── One-Time Setup ───
    doc.add_heading('One-Time Cluster Setup (Complete Commands)', level=1)
    doc.add_paragraph(
        'All the setup below is one-time per cluster. Same steps apply to EKS/AKS/GKE. '
        'After this, every new deployment is just a git push.'
    )

    # Step 1: Create Cluster
    doc.add_heading('Step 1: Create the Cluster', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'export KIND_EXPERIMENTAL_PROVIDER=podman\n'
        'kind create cluster --name java-demo-cluster --config kind-cluster-config.yaml\n'
        'kubectl get nodes\n'
        '# Expected: 3 nodes (1 control-plane + 2 workers)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 2: Install ArgoCD
    doc.add_heading('Step 2: Install ArgoCD', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'kubectl create namespace argocd\n'
        'kubectl apply -n argocd -f https://raw.githubusercontent.com/argoproj/argo-cd/stable/manifests/install.yaml\n'
        'kubectl wait --for=condition=available --timeout=300s deployment/argocd-server -n argocd\n'
        '# Get admin password:\n'
        'kubectl get secret argocd-initial-admin-secret -n argocd -o jsonpath="{.data.password}" | base64 -d'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 3: Install Image Updater
    doc.add_heading('Step 3: Install ArgoCD Image Updater', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'kubectl apply -n argocd \\\n'
        '  -f https://raw.githubusercontent.com/argoproj-labs/argocd-image-updater/v0.14.0/manifests/install.yaml\n'
        'kubectl get pods -n argocd -l app.kubernetes.io/name=argocd-image-updater'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 4: Install Kyverno
    doc.add_heading('Step 4: Install Kyverno', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'kubectl create -f https://github.com/kyverno/kyverno/releases/download/v1.12.0/install.yaml\n'
        'kubectl wait --for=condition=available --timeout=300s deployment/kyverno-admission-controller -n kyverno'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 5: Cosign Keys
    doc.add_heading('Step 5: Generate Cosign Key Pair', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'cosign generate-key-pair\n'
        '# Creates: cosign.key (PRIVATE → GitHub Secrets) and cosign.pub (PUBLIC → Kyverno policy)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 6: Create Secrets
    doc.add_heading('Step 6: Create All Credential Secrets', level=2)
    doc.add_paragraph('ArgoCD namespace (for Image Updater):')
    p = doc.add_paragraph()
    run = p.add_run(
        '# Git repo credentials (for Image Updater git write-back)\n'
        'kubectl create secret generic repo-crd-poc -n argocd \\\n'
        '  --from-literal=type=git \\\n'
        '  --from-literal=url=https://github.com/<org>/<repo>.git \\\n'
        '  --from-literal=username=<github-user> \\\n'
        '  --from-literal=password=<github-pat>\n'
        'kubectl label secret repo-crd-poc -n argocd argocd.argoproj.io/secret-type=repository\n'
        '\n'
        '# Docker Hub credentials (for Image Updater to poll tags)\n'
        'kubectl create secret docker-registry dockerhub-creds -n argocd \\\n'
        '  --docker-server=https://registry-1.docker.io \\\n'
        '  --docker-username=<docker-user> \\\n'
        '  --docker-password=<docker-password>'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('Kyverno namespace (for signature verification):')
    p = doc.add_paragraph()
    run = p.add_run(
        '# Two secrets needed (Docker Hub uses two registry URLs)\n'
        'kubectl create secret docker-registry docker-hub-creds -n kyverno \\\n'
        '  --docker-server=https://index.docker.io/v1/ \\\n'
        '  --docker-username=<docker-user> --docker-password=<docker-password>\n'
        'kubectl annotate secret docker-hub-creds -n kyverno kyverno.io/docker-config=true\n'
        '\n'
        'kubectl create secret docker-registry docker-hub-creds-v2 -n kyverno \\\n'
        '  --docker-server=https://registry-1.docker.io \\\n'
        '  --docker-username=<docker-user> --docker-password=<docker-password>\n'
        'kubectl annotate secret docker-hub-creds-v2 -n kyverno kyverno.io/docker-config=true'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('GitHub Secrets (in GitHub UI → Settings → Secrets → Actions):')
    gh_secrets = [
        ('Secret', 'Value'),
        ('DOCKER_USERNAME', 'Docker Hub username'),
        ('DOCKER_PASSWORD', 'Docker Hub access token'),
        ('COSIGN_PRIVATE_KEY', 'Contents of cosign.key file'),
        ('COSIGN_PASSWORD', 'Cosign key password'),
    ]
    table = doc.add_table(rows=len(gh_secrets), cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(gh_secrets):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Step 7: Deploy
    doc.add_heading('Step 7: Deploy the Application', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'kubectl apply -f k8s/platform/kyverno-verify-images.yaml    # Signature verification policy\n'
        'kubectl apply -f k8s/platform/argocd-application-v2.yaml    # ArgoCD app with Image Updater annotations'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Step 8: Verify
    doc.add_heading('Step 8: Verify the Setup', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'kubectl get pods -n argocd        # ArgoCD + Image Updater\n'
        'kubectl get pods -n kyverno       # Kyverno\n'
        'kubectl get pods -n default       # App (3 replicas)\n'
        'kubectl port-forward svc/java-demo-app 8083:80 -n default &\n'
        'curl http://localhost:8083/'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Setup summary table
    doc.add_paragraph()
    setup_summary = [
        ('Component', 'Namespace', 'Purpose'),
        ('ArgoCD v3.3.6', 'argocd', 'GitOps controller — syncs K8s manifests from Git'),
        ('Image Updater v0.14.0', 'argocd', 'Watches Docker Hub → commits new tags to Git'),
        ('Kyverno v1.12.0', 'kyverno', 'Admission controller — blocks unsigned images'),
        ('Cosign keys', 'GitHub Secrets + Kyverno', 'Signs images in CI, verifies in cluster'),
        ('App (3 replicas)', 'default', 'Spring Boot API with ClusterIP service'),
    ]
    table = doc.add_table(rows=len(setup_summary), cols=3, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(setup_summary):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('After this one-time setup, every new deployment is just: git push')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_page_break()

    # ─── Summary ───
    doc.add_heading('Summary', level=1)
    summary = [
        ('', 'Approach 1 (Demo)', 'Approach 2 (Production)'),
        ('Pipeline', '4 jobs, writes to Git', '2 jobs, read-only'),
        ('Security', 'Trivy scan', 'Trivy + Cosign + Kyverno'),
        ('Deployment', 'Pipeline pushes manifest', 'Image Updater auto-detects'),
        ('Complexity', 'More pipeline code', 'More K8s components'),
        ('Recommendation', 'Learning & demos', 'Production & client environments'),
    ]
    table = doc.add_table(rows=len(summary), cols=3, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(summary):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Both approaches achieve the same goal: code push → automatic deployment. ')
    p.add_run('Approach 2 adds image signing, admission control, and cleaner separation of concerns — '
              'making it the right choice for production and client-facing environments.')

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'CI-CD-Demo-Presentation.docx')
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    img1 = generate_approach1()
    img2 = generate_approach2()
    img3 = generate_comparison()
    docx_path = build_docx(img1, img2, img3)
    print(f"\n🎉 All done! Open: {docx_path}")
